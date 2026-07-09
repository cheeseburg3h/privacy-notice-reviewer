from __future__ import annotations

from pathlib import Path


def render_docx_from_markdown(markdown_path: str | Path, output_path: str | Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("python-docx is required for optional DOCX rendering.") from exc

    document = Document()
    for line in Path(markdown_path).read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output

